from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "实验13_121072021030_林立洲.docx"
IMG = ROOT / "report-images"


def set_run_font(run, size=11, name="宋体", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True


def add_center_title(doc: Document, text: str, size=20) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, name="黑体", bold=True)


def add_heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 13, name="黑体", bold=True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, 11, "宋体")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, 11, "宋体")


def add_code(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, 9, "Consolas")


def add_picture(doc: Document, image_name: str, caption: str, width=6.2) -> None:
    path = IMG / image_name
    if not path.exists():
        add_para(doc, f"{caption}：图片文件暂未生成，详见对应文本输出文件。")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, 10, "宋体")


def read_text(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def extract_block(path: str, marker: str) -> str:
    lines = read_text(path).splitlines()
    start = next((i for i, line in enumerate(lines) if marker in line), None)
    if start is None:
        return ""
    block = []
    depth = 0
    seen_open = False
    for line in lines[start:]:
        block.append(line)
        depth += line.count("{")
        if "{" in line:
            seen_open = True
        depth -= line.count("}")
        if seen_open and depth == 0:
            break
    return "\n".join(block)


def add_cover(doc: Document) -> None:
    add_center_title(doc, "福建师范大学", 20)
    add_center_title(doc, "《软件体系结构与设计模式》实验报告", 18)
    doc.add_paragraph()
    add_center_title(doc, "实验13 模板方法与访问者模式", 18)
    doc.add_paragraph()

    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    data = [
        ("课程名称", "软件体系结构与设计模式"),
        ("年级专业", "2023级软件工程"),
        ("班级", "软工2班"),
        ("学号", "121072021030"),
        ("姓名", "林立洲"),
        ("实验环境", "Windows、Enterprise Architect 12、IntelliJ IDEA、Java、Maven"),
        ("实验内容", "模板方法模式银行利息计算；访问者模式 OA 员工信息管理"),
    ]
    for row, (k, v) in zip(table.rows, data):
        for idx, text in enumerate([k, v]):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, 11, "宋体", bold=(idx == 0))
    doc.add_page_break()


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    setup_styles(doc)

    add_cover(doc)

    add_heading(doc, "一、实验目的", 1)
    add_para(doc, "本次实验围绕模板方法模式和访问者模式展开，目标是在具体业务场景中理解设计模式的结构、适用条件和实现方式。通过银行利息计算问题，可以体会父类统一控制算法骨架、子类实现可变步骤的思想；通过公司 OA 员工信息管理问题，可以体会将不同部门对同一批对象的操作封装为访问者的思想。")
    add_bullets(doc, [
        "理解模板方法模式中固定流程与可变步骤分离的设计方法。",
        "理解访问者模式中对象结构与外部操作分离的设计方法。",
        "掌握使用 UML 类图和顺序图表达类职责、对象交互和模式结构。",
        "掌握使用 Java 与 Maven 实现、运行和测试两个设计模式案例。",
    ])

    add_heading(doc, "二、实验内容与需求分析", 1)
    add_heading(doc, "2.1 模板方法模式银行利息计算", 2)
    add_para(doc, "银行利息计算流程包括根据账号查询用户信息、判断账户类型、根据账户类型计算利息以及显示利息结果。活期账户和定期账户虽然计算公式不同，但业务处理步骤具有明显的固定顺序，因此适合使用模板方法模式。系统将公共流程放在抽象模板类中，将具体利息计算交给子类实现，客户端只需要调用统一的模板方法即可完成计算。")
    add_para(doc, "本实验中，活期账户使用“余额 × 年利率 × 存期月数 / 12”的公式；定期账户使用同一基础公式，并在存期满 12 个月时追加 5% 奖励利息。这样既保持流程统一，又体现不同账户类型在可变步骤上的差异。")

    add_heading(doc, "2.2 访问者模式员工信息管理", 2)
    add_para(doc, "公司 OA 员工信息管理子系统中的员工类型包括正式员工和临时工。正式员工每周标准工时为 40 小时，超过部分按 100 元/小时计算加班费，少于 40 小时按 80 元/小时扣除请假工资；临时工工资按实际工时和小时工资计算。人力资源部关注工时、加班和请假情况，财务部关注工资计算，两类操作作用于同一批员工对象。")
    add_para(doc, "由于员工类型相对稳定，而部门统计需求可能变化，如果把所有统计和计算逻辑都写入员工类，会导致员工类职责过重。访问者模式将 HR 工时汇总和 Finance 工资计算分别封装到访问者类中，员工对象只负责接收访问者并回调相应的 visit 方法，从而提高职责划分的清晰度。")

    add_heading(doc, "三、系统设计", 1)
    add_heading(doc, "3.1 模板方法模式设计", 2)
    add_para(doc, "模板方法模式部分以 InterestCalculatorTemplate 作为抽象模板类，其中 repository 用于访问账户数据，calculateInterest(accountNo) 是模板方法，并使用 final 修饰，防止子类改变整体流程。queryAccount、checkAccount 和 displayInterest 是公共步骤；doCalculateInterest 是抽象步骤，由活期账户和定期账户计算器实现。")
    add_para(doc, "BankAccount 表示账户实体，保存账号、户名、余额、年利率、存期和账户类型。AccountRepository 用 Map 模拟账户数据库，提供添加账户和按账号查询账户的功能。BankInterestDemo 作为演示类，创建账户仓库和两个具体计算器，分别完成 C001 活期账户和 S001 定期账户的利息计算。")
    add_picture(doc, "TemplateMethod银行利息计算类图.png", "图3-1 TemplateMethod 银行利息计算类图", 6.4)
    add_picture(doc, "TemplateMethod银行利息计算顺序图.png", "图3-2 TemplateMethod 银行利息计算顺序图", 6.4)

    add_heading(doc, "3.2 访问者模式设计", 2)
    add_para(doc, "访问者模式部分以 Employee 作为抽象元素接口，FullTimeEmployee 和 TemporaryEmployee 作为具体元素。DepartmentVisitor 是抽象访问者，声明访问正式员工和临时工的两个 visit 方法。HRDepartmentVisitor 汇总工时、加班和请假情况，FinanceDepartmentVisitor 计算工资。EmployeeList 是对象结构，保存员工集合，并在 accept(visitor) 中依次让每个员工接收访问者。")
    add_para(doc, "这种设计使员工类中不出现 HR 和 Finance 的具体统计逻辑，员工对象只暴露必要数据并实现 accept 方法。当新增一个部门统计需求时，可以新增访问者类，而不需要修改 FullTimeEmployee、TemporaryEmployee 和 EmployeeList 的基本结构。")
    add_picture(doc, "Visitor员工信息管理类图.png", "图3-3 Visitor 员工信息管理类图", 6.4)
    add_picture(doc, "Visitor员工信息管理顺序图.png", "图3-4 Visitor 员工信息管理顺序图", 6.4)

    add_heading(doc, "四、系统实现", 1)
    add_heading(doc, "4.1 项目结构", 2)
    add_para(doc, "项目采用 Maven 管理，源代码位于 src/main/java/com/lzl/experiment13。template 子包实现模板方法模式银行利息计算，visitor 子包实现访问者模式员工信息管理，MainApp 负责统一启动演示和测试。")
    add_picture(doc, "IntelliJ项目结构目录图.png", "图4-1 项目结构目录图", 6.2)

    add_heading(doc, "4.2 模板方法模式核心代码", 2)
    add_para(doc, "InterestCalculatorTemplate 的 calculateInterest 方法是整个利息计算流程的骨架。客户端传入账号后，方法先查询账户，再检查账户是否存在，随后调用 doCalculateInterest 执行子类差异步骤，最后显示利息结果。")
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/template/InterestCalculatorTemplate.java", "public final void calculateInterest"))
    add_para(doc, "活期账户计算器只实现具体公式，不改变父类定义的流程。")
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/template/CurrentAccountInterestCalculator.java", "public double doCalculateInterest"))
    add_para(doc, "定期账户计算器在基础公式上增加满 12 个月奖励利息规则，体现同一模板流程下不同子类的业务差异。")
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/template/SavingAccountInterestCalculator.java", "public double doCalculateInterest"))
    add_para(doc, "AccountRepository 使用 LinkedHashMap 保存账户数据，findByAccountNo 用于模拟根据账号查询用户信息。")
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/template/AccountRepository.java", "public BankAccount findByAccountNo"))

    add_heading(doc, "4.3 访问者模式核心代码", 2)
    add_para(doc, "Employee 接口要求所有员工对象都能接收部门访问者，正式员工和临时工在 accept 方法中分别调用 visitor.visit(this)，从而形成访问者模式中的双分派结构。")
    add_code(doc, read_text("src/main/java/com/lzl/experiment13/visitor/Employee.java"))
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/visitor/FullTimeEmployee.java", "public void accept"))
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/visitor/TemporaryEmployee.java", "public void accept"))
    add_para(doc, "DepartmentVisitor 声明访问两类员工的方法，使不同部门访问者都能对正式员工和临时工执行各自逻辑。")
    add_code(doc, read_text("src/main/java/com/lzl/experiment13/visitor/DepartmentVisitor.java"))
    add_para(doc, "HRDepartmentVisitor 负责统计工时、加班和请假；FinanceDepartmentVisitor 负责根据制度计算工资。两者访问同一 EmployeeList，但输出的报告内容不同。")
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/visitor/HRDepartmentVisitor.java", "public void visit(FullTimeEmployee"))
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/visitor/FinanceDepartmentVisitor.java", "public void visit(FullTimeEmployee"))
    add_para(doc, "EmployeeList 作为对象结构，保存员工集合并统一触发访问。")
    add_code(doc, extract_block("src/main/java/com/lzl/experiment13/visitor/EmployeeList.java", "public void accept"))

    add_heading(doc, "五、测试与运行结果", 1)
    add_heading(doc, "5.1 Maven 构建结果", 2)
    add_para(doc, "在项目根目录执行 mvn clean package，Maven 完成清理、编译、打包流程，并生成 target/experiment13-template-visitor-1.0.0.jar。构建命令的完整输出保存于 maven-package-result.txt。")
    add_picture(doc, "Maven构建成功结果.png", "图5-1 Maven 构建结果", 6.2)

    add_heading(doc, "5.2 程序运行结果", 2)
    add_para(doc, "执行 java -cp target/experiment13-template-visitor-1.0.0.jar com.lzl.experiment13.MainApp 后，程序依次运行模板方法模式演示、访问者模式演示和控制台测试。完整输出保存于 run-result.txt。")
    add_picture(doc, "MainApp运行结果.png", "图5-2 MainApp 完整运行结果", 6.2)
    add_para(doc, "模板方法模式运行结果表明，活期账户 C001 的利息为 17.50 元，定期账户 S001 按满 12 个月奖励规则得到 378.00 元。两个账户都遵循查询、检查、计算、显示的统一流程。")
    add_picture(doc, "模板方法模式运行结果.png", "图5-3 模板方法模式运行结果", 6.2)
    add_para(doc, "访问者模式运行结果表明，HR 访问者输出正式员工工作小时数、临时工实际工作小时数、加班小时数、请假小时数和总工时；Finance 访问者输出正式员工与临时工的工资计算过程以及总工资。")
    add_picture(doc, "访问者模式运行结果.png", "图5-4 访问者模式运行结果", 6.2)

    add_heading(doc, "5.3 测试代码结果", 2)
    add_para(doc, "Experiment13TestRunner 提供控制台测试，覆盖活期账户利息、定期账户利息、不存在账号提示、模板方法完整流程、正式员工加班与请假时长、临时工工资、财务部总工资以及两个访问者报告差异。运行结果显示所有测试通过。")

    add_heading(doc, "六、实验总结", 1)
    add_para(doc, "通过本次实验，我对模板方法模式中固定流程和可变步骤的分离有了更清晰的理解。银行利息计算中，父类统一控制查询、检查、计算和显示流程，子类只负责具体计算公式，因此业务流程不容易被子类随意改动，程序结构也更加稳定。")
    add_para(doc, "通过访问者模式的实现，我理解了如何将不同部门对员工数据的操作封装到访问者中。同一批员工对象可以被 HR 和 Finance 两个访问者以不同方式处理，员工类本身不需要承担工时汇总和工资计算的全部职责。这样的设计在对象结构稳定、操作类型可能扩展的场景中比较合适。")
    add_para(doc, "在完成 UML 建模、Java 编码、Maven 构建和控制台测试的过程中，我进一步体会到先设计再实现的重要性。类图帮助我明确了类之间的继承、聚合、依赖和接口关系，顺序图帮助我梳理了运行时对象之间的调用顺序。整体来看，本实验加深了我对模板方法模式和访问者模式适用场景、优点和实现细节的理解。")

    add_heading(doc, "七、附录", 1)
    add_heading(doc, "7.1 项目目录与文件", 2)
    add_bullets(doc, [
        "Maven 配置文件：pom.xml",
        "主程序入口：src/main/java/com/lzl/experiment13/MainApp.java",
        "模板方法模式代码：src/main/java/com/lzl/experiment13/template/",
        "访问者模式代码：src/main/java/com/lzl/experiment13/visitor/",
        "UML 图源：uml/",
        "UML 与运行结果图片：report-images/",
        "运行输出：run-result.txt",
        "Maven 构建输出：maven-package-result.txt",
        "EA 项目文件：实验13_121072021030_林立洲.eap",
    ])
    add_heading(doc, "7.2 运行命令", 2)
    add_code(doc, "mvn clean package\njava -cp target/experiment13-template-visitor-1.0.0.jar com.lzl.experiment13.MainApp")
    add_heading(doc, "7.3 Git 操作结果", 2)
    if (IMG / "Git提交成功结果.png").exists():
        add_picture(doc, "Git提交成功结果.png", "图7-1 Git 操作结果", 6.2)
    else:
        add_para(doc, "Git 操作结果保存于 git-result.txt。")
    add_para(doc, "本实验未填写 GitHub 或 Gitee 远程仓库地址，仅保留本地 Git 仓库和本地提交记录。")

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
